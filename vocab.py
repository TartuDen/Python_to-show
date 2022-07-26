import xlrd
from docx import Document
import sys
import random
from tkinter import *
from tkinter import ttk
import sqlite3
from tkinter import messagebox
from googletrans import Translator
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from PyDictionary import PyDictionary
from tkinter import colorchooser
import math
import time
import os
import pyautogui as pg
import sys
import codecs
#
# search mayusculas y menusculas -DONE
# check the new table and compare with the old one -DONE
# possibility to add new lines - in progress


class Create_check_db(object):
    def __init__(self):

        # self.Create()  # - update DB
        # self.Check_Combine_db()
        pass

    def Create(self):
        conn = sqlite3.connect('V3.db')
        c = conn.cursor()

        c.execute("""CREATE TABLE voc_table(
        id_ integer,
        rus_ text,
        eng_ text,
        esp_ text)""")
        idx = 0

        wordDoc = Document('Vocabulario_Completo_Ruso_Eng_Esp_Eesti.docx')
        for table in wordDoc.tables:
            for row in table.rows:
                list_ = []

                for cell in row.cells:
                    temp_ = ''
                    for ch in cell.text:
                        if ord(ch) in range(65535):
                            temp_ = temp_ + ch
                    try:
                        list_.append(temp_)
                    except:
                        None
                if len(list_) == 3:
                    c.execute("INSERT INTO voc_table VALUES (?,?,?,?,?)", (idx, list_[0], list_[1], list_[2],''))
                    idx += 1
                elif len(list_) == 2:
                    c.execute("INSERT INTO voc_table VALUES (?,?,?,?,?)", (idx, list_[0], list_[1], 'None',''))
                    idx += 1
                conn.commit()

    def Check_Combine_db(self):
        conn = sqlite3.connect('V1.db')
        c = conn.cursor()
        conn2 = sqlite3.connect('V3.db')
        c2 = conn2.cursor()
        c.execute("SELECT* FROM voc_table")
        c2.execute("SELECT* FROM voc_table")
        list_to1 = []
        list_to2 = []
        total = []
        for st in c.fetchall():
            list_to1.append(st)

        for st in c2.fetchall():
            list_to2.append(st)
        for ln2 in list_to2:
            match = False
            for ln1 in list_to1:
                if ln1[1] == ln2[1] and ln1[2] == ln2[2]:
                    match = True
                    break
            if match == False:
                total.append(ln2)

        #print('Lines were changed: ', total)
        for st in total:
            if len(st) == 4:
                c.execute("INSERT INTO voc_table VALUES(?,?,?,?,?)", (st[0], st[1], st[2], st[3],''))
            elif len(st) == 3:
                c.execute("INSERT INTO voc_table VALUES(?,?,?,?,?)", (st[0], st[1], st[2], 'To be added',''))
        conn.commit()
        conn2.commit()


class Translate_selected(object):
    def __init__(self, graph_):
        self.graph_ = graph_
        self.translator = Translator()
        self.vocab = PyDictionary()

    def method_rus_esp(self):
        selected = self.graph_.text_rus.get(SEL_FIRST, SEL_LAST)
        transl_from_rus_esp = self.translator.translate(selected, dest='es')

        is_None = self.graph_.text_esp.get('1.0', 'end')
        transl_from_rus_eng = self.translator.translate(selected)

        if transl_from_rus_esp.text not in self.graph_.text_esp.get('1.0', 'end') and transl_from_rus_esp.text not in (self.graph_.text_esp.get('1.0', 'end').title()):
            if 'None' in is_None:
                self.graph_.text_esp.delete('1.0', 'end')
                self.graph_.text_esp.insert('end', '%s  -  %s' % (selected, transl_from_rus_esp.text))
            elif len(is_None)<2:
                self.graph_.text_esp.insert('end', '%s  -  %s' % (selected, transl_from_rus_esp.text))
            else:
                self.graph_.text_esp.insert('end', '\n%s  -  %s' % (selected, transl_from_rus_esp.text))

        if (transl_from_rus_eng.text) not in self.graph_.text_eng.get('1.0', 'end'):
            self.graph_.text_eng.insert('end', '\n%s' % transl_from_rus_eng.text)

    def method_eng_esp_and_meaning_esp(self):
        try:
            selected = self.graph_.text_eng.get(SEL_FIRST, SEL_LAST)
            transl_from_eng_ee = self.translator.translate(selected, dest='ee')
            if transl_from_eng_ee==None:
                transl_from_eng_ee = self.vocab.translate(selected, dest='ee')

            self.graph_.text_rus.insert('end', '\n%s\n' % ('-' * 15))
            self.graph_.text_rus.insert('end', '%s\n' % ('Eesti:'))
            self.graph_.text_rus.insert('end', '%s  -  %s' % (selected, transl_from_eng_ee.text))

        except:
            pass
        try:
            selected = self.graph_.text_esp.get(SEL_FIRST, SEL_LAST)
        except:
            pass

        try:

            transl_from_eng_esp = self.translator.translate(selected, dest='es')
        except:
            transl_from_eng_esp=('doesnt work yet')

        is_None = self.graph_.text_esp.get('1.0', 'end')
        if transl_from_eng_esp!='doesnt work yet' and transl_from_eng_esp.text not in self.graph_.text_esp.get('1.0', 'end') and transl_from_eng_esp.text not in (self.graph_.text_esp.get('1.0', 'end').title()):
            if 'None' in is_None or len(is_None)<2:
                self.graph_.text_esp.delete('1.0', 'end')
                self.graph_.text_esp.insert('end', '%s - %s' % (selected, transl_from_eng_esp.text))
            else:
                self.graph_.text_esp.insert('end', '\n%s - %s' % (selected, transl_from_eng_esp.text))
        meaning = self.vocab.meaning(selected)
        try:
            self.graph_.text_eng.insert('end', '\n\n')
            self.graph_.text_eng.insert('end', '--------------')
            self.graph_.text_eng.insert('end', '\n')
            self.graph_.text_eng.insert('end', 'Noun:  ʻ%sʼ' % selected.upper())
            self.graph_.text_eng.insert('end', '\n')
            for sent_ in meaning['Noun']:
                self.graph_.text_eng.insert('end', '* ' + sent_)
                self.graph_.text_eng.insert('end', '\n')

            self.graph_.text_eng.insert('end', '\n\n')
            self.graph_.text_eng.insert('end', '--------------')
            self.graph_.text_eng.insert('end', '\n')
            self.graph_.text_eng.insert('end', 'Verb:  ʻ%sʼ' % selected.upper())
            self.graph_.text_eng.insert('end', '\n')
            for sent_ in meaning['Verb']:
                self.graph_.text_eng.insert('end', '* ' + sent_)
                self.graph_.text_eng.insert('end', '\n')
            self.graph_.text_eng.insert('end', '\n\n')
            self.graph_.text_eng.insert('end', '--------------')
            self.graph_.text_eng.insert('end', '\n')
            self.graph_.text_eng.insert('end', 'Adjective:  ʻ%sʼ' % selected.upper())
            self.graph_.text_eng.insert('end', '\n')
            for sent_ in meaning['Adjective']:
                self.graph_.text_eng.insert('end', '* ' + sent_)
                self.graph_.text_eng.insert('end', '\n')
        except:
            pass
        if self.graph_.get_BrowserGoogle_var.get() == 'True':
            self.getBrowserGoogle(selected)
        if self.graph_.get_BrowserWR_var.get() == 'True':
            self.getBrowserWR(selected)

    def getBrowserGoogle(self, selected):
        driver = webdriver.Chrome()
        driver.get(r'https://translate.google.com/?hl=es')
        text_form = driver.find_element_by_id('source')
        text_form.send_keys(selected)
        text_form.send_keys(Keys.RETURN)

    def getBrowserWR(self, selected):
        driver = webdriver.Chrome()
        driver.get('http://www.wordreference.com/definicion')
        element_enter_text = driver.find_element_by_id('si')
        element_enter_text.send_keys(selected)
        pg.hotkey('enter')


class Graph(object):
    def __init__(self, master):
        start_graph=time.process_time()

        self.color = 0
        self.sel = StringVar()
        self.fav_sel = StringVar()
        self.eesti_sel = StringVar()
        self.goBack = StringVar()
        self.Color_select = StringVar()
        self.get_BrowserGoogle_var = StringVar()
        self.get_BrowserWR_var = StringVar()

        self.master = master
        self.master.config(bg='LightCyan3')
        self.text_rus = Text(self.master, width=25, height=15, wrap='word', bg='old lace')
        self.text_rus.place(x=10, y=120)

        self.scroll_rus=ttk.Scrollbar(self.text_rus,orient=VERTICAL, command=self.text_rus.yview)
        self.scroll_rus.place(relx=0.95,rely=0.1,width=10,height=200)
        self.text_rus.config(yscrollcommand=self.scroll_rus.set)

        self.text_eng = Text(self.master, width=25, height=15, wrap='word', bg='old lace')
        self.text_eng.place(x=240, y=120)

        self.text_esp = Text(self.master, width=25, height=15, wrap='word', bg='old lace')
        self.text_esp.place(x=470, y=120)



        self.word_ = W0rd(self)
        self.f_word = Favorits(self)
        self.edit_ = Edit_(self)
        self.sel_txt = Translate_selected(self)
        self.eesti_ = Eesti(self)
        self.nuevo_ = Nuevo(self)
        self.borrar_=Delete_row(self)
        self.fonts_=Fonts(self)

        self.master.protocol("WM_DELETE_WINDOW", self.f_word.on_closing)
        self.master.bind('<Escape>', lambda e: self.f_word.on_closing())

        self.b_continue = Button(self.master, text='Continue', command=self.word_.getWord)
        self.b_continue.place(relx=0.01, y=400, width=150)

        #self.text_rus.bind('<Double-Button-1>', lambda e: self.word_.getWord())

        # self.color_combo = ttk.Combobox(self.master, textvariable=self.Color_select, values=('Blue', 'Green', 'Red', 'White'))
        # self.color_combo.place(relx=0.01, rely=0.01)
        # self.Color_select.set('Blue')

        # self.b_back = Button(self.master, text='<---- Back', command=self.backwards)
        # self.b_back.place(relx=0.39, y=400, width=150)

        self.master.bind('<F1>', lambda e: self.sel_txt.method_rus_esp())
        self.master.bind('<F2>', lambda e: self.sel_txt.method_eng_esp_and_meaning_esp())
        self.master.bind('<F3>', lambda e: self.edit_.edit_word_file())

        self.label_info1 = ttk.Label(self.master)
        self.label_info1.place(relx=0.25, y=2, width=200, height=67)
        self.label_info1.config(text=' F1 - rus->esp;\n F2 - eng->esp + eng meaning;\n F2 + Google;\n F3 - word doc search;', relief=RIDGE, background='#C5D8D8')

        self.label_info2 = ttk.Label(self.master)
        self.label_info2.place(relx=0.54, y=2, width=100, height=50)
        self.label_info2.config(text='', relief=RIDGE, background='#C5D8D8')

        self.label_choice_buttons=ttk.Label(self.master, relief=RIDGE)
        self.label_choice_buttons.place(x=9,y=1,width=80,height=118)

        self.b_search = Button(self.master, text=' (¬_¬) <SEARCH>  ಠ_ಠ ', command=self.search_)
        self.b_search.place(relx=0.77, y=400, width=150)

        # self.b_color = Button(self.master, text='_color_', command=self.color_apply)
        # self.b_color.place(relx=0.35, y=400, width=150)

        self.Back_checkbox = ttk.Checkbutton(self.label_choice_buttons, text='GoBack')
        self.Back_checkbox.place(x=3, y=1)
        self.Back_checkbox.config(variable=self.goBack, onvalue='True', offvalue='False')

        self.Frase_button = ttk.Checkbutton(self.label_choice_buttons, text='Frases')
        self.Frase_button.place(x=3, y=21)
        self.Frase_button.config(variable=self.sel, onvalue='True', offvalue='False')

        self.Fav_button = ttk.Checkbutton(self.label_choice_buttons, text='Favoritos')
        self.Fav_button.place(x=3, y=101)
        self.Fav_button.config(variable=self.fav_sel, onvalue='True', offvalue='False')

        self.Eesti_button = ttk.Checkbutton(self.label_choice_buttons, text='Eesti')
        self.Eesti_button.place(x=3, y=81)
        self.Eesti_button.config(variable=self.eesti_sel, onvalue='True', offvalue='False')

        self.browsG_button = ttk.Checkbutton(self.label_choice_buttons, text='Google')
        self.browsG_button.place(x=3, y=41)
        self.browsG_button.config(variable=self.get_BrowserGoogle_var, onvalue='True', offvalue='False')

        self.browsWR_button = ttk.Checkbutton(self.label_choice_buttons, text='WordRef')
        self.browsWR_button.place(x=3, y=61)
        self.browsWR_button.config(variable=self.get_BrowserWR_var, onvalue='True', offvalue='False')

        self.b_fav_add = Button(self.master, text=' Añadir ', command=self.f_word.add_)
        self.b_fav_add.place(relx=0.75, y=60, width=80)

        self.b_limpiar = Button(self.master, text=' Limpiar ', command=self.nuevo_.clear)
        self.b_limpiar.place(relx=0.75, y=30, width=80)

        self.b_crear = Button(self.master, text=' Crear nuevo ', command=self.nuevo_.crear_)
        self.b_crear.place(relx=0.87, y=30, width=80)

        self.b_fav_rem = Button(self.master, text=' Quitar ', command=self.f_word.rem_)
        self.b_fav_rem.place(relx=0.87, y=60, width=80)

        self.b_delete = Button(self.master, text=' Borrar ', command=self.borrar_.toDelete)
        self.b_delete.place(relx=0.87, y=89, width=80)

        self.edit_esp_button = Button(self.text_esp, text=' editar ', command=self.edit_.generic_func)
        self.edit_esp_button.place(relx=0.9, x=-30, rely=0.94, y=-10, width=50)

        self.font_frame=Frame(self.master, bg='grey')
        self.font_frame.place(relx=0.015, y=365, width=85, height=20)



        self.bold_font_button=Button(self.font_frame,text='B',bg='#C5D8D8',command=self.fonts_.bold_)
        self.bold_font_button.place(x=2,rely=0.03,width=20, height=19)

        self.italic_font_button=Button(self.font_frame,text='It',bg='#C5D8D8',command=self.fonts_.italic_)
        self.italic_font_button.place(x=23,rely=0.03,width=20, height=19)

        self.color_button=Button(self.font_frame,text='Cc',bg='#46a3ff',command=self.fonts_.blue_bg)
        self.color_button.place(x=44,rely=0.03,width=20, height=19)

        self.clear_tag_button=Button(self.font_frame,text='tg',bg='white',command=self.fonts_.clear_tag)
        self.clear_tag_button.place(x=65,rely=0.03,width=20, height=19)
        stop_graph=time.process_time()
        print('Graph time: %f'%(stop_graph- start_graph))



    def search_(self):
        idx_srch = []
        s_ch = self.text_rus.get('1.0', 'end')
        s_ch = s_ch.strip()
        for words in self.word_.search_list_rus:
            if 'Frases' not in words:
                if (s_ch in words) or (s_ch.title() in words):
                    idx = self.word_.search_list_rus.index(words)
                    idx_srch.append(idx)
        for words in self.word_.search_list_eng:
            if (s_ch in words) or (s_ch.title() in words):
                idx = self.word_.search_list_eng.index(words)
                idx_srch.append(idx)

        for words in self.word_.search_list_esp:
            if (s_ch in words) or (s_ch.title() in words):
                idx = self.word_.search_list_esp.index(words)
                idx_srch.append(idx)

        try:
            self.text_rus.delete('1.0', 'end')
            self.text_eng.delete('1.0', 'end')
            self.text_esp.delete('1.0', 'end')
        except:
            None
        count = 1
        test_rus=[]

        for i in idx_srch:
            test_rus.append(self.word_.search_list_rus[i])

        test_rus.sort(key = lambda s: len(s),reverse=True)




        if 'Eesti:' in test_rus[0]:
            for w in test_rus:
                self.text_rus.insert(INSERT, '%d)  ' % count + w)
                self.text_rus.insert(INSERT, '\n ===============\n')
                count += 1
        else:
            for i in idx_srch:
                self.text_rus.insert(INSERT, '%d)  ' % count + self.word_.search_list_rus[i])
                self.text_rus.insert(INSERT, '\n ===============\n')

                self.text_eng.insert(INSERT, '%d)  ' % count + self.word_.search_list_eng[i])
                self.text_eng.insert(INSERT, '\n ===============\n')
                self.text_esp.insert(INSERT, '%d)  ' % count + self.word_.search_list_esp[i])
                self.text_esp.insert(INSERT, '\n ===============\n')
                count += 1


    # def color_apply(self):
    #     #'Blue' - '#46a3ff','Green' - '#00c100','Red' - '#ff4242','White' - '#ffffff'
    #     # self.color = colorchooser.askcolor(initialcolor='#FFFFFF')  # this opens the color table
    #     if self.Color_select.get() == 'Blue':

    #         try:
    #             sel_first_index = (self.text_rus.index(SEL_FIRST))
    #             self.text_rus.insert(sel_first_index, 'ʻ')
    #             sel_last_index = (self.text_rus.index(SEL_LAST))
    #             self.text_rus.insert(sel_last_index, 'ʼ')
    #             self.text_rus.tag_add('blue', sel_first_index, sel_last_index)
    #         except:
    #             None
    #         try:
    #             sel_first_index = (self.text_eng.index(SEL_FIRST))
    #             self.text_eng.insert(sel_first_index, 'ʻ')
    #             sel_last_index = (self.text_eng.index(SEL_LAST))
    #             self.text_eng.insert(sel_last_index, 'ʼ')
    #             self.text_eng.tag_add('blue', sel_first_index, sel_last_index)
    #         except:
    #             None
    #         try:
    #             sel_first_index = (self.text_esp.index(SEL_FIRST))
    #             self.text_esp.insert(sel_first_index, 'ʻ')
    #             sel_last_index = (self.text_esp.index(SEL_LAST))
    #             self.text_esp.insert(sel_last_index, 'ʼ')
    #             self.text_esp.tag_add('blue', sel_first_index, sel_last_index)
    #         except:
    #             None

    #     elif self.Color_select.get() == 'Green':
    #         sel_first_index = (self.text_rus.index(SEL_FIRST))
    #         self.text_rus.insert(sel_first_index, '´')
    #         sel_last_index = (self.text_rus.index(SEL_LAST))
    #         self.text_rus.insert(sel_last_index, '´')
    #         self.text_rus.tag_add('green', sel_first_index, sel_last_index)

    #     self.text_rus.tag_config('blue', background='lightblue', foreground='black')
    #     self.text_eng.tag_config('blue', background='lightblue', foreground='black')
    #     self.text_esp.tag_config('blue', background='lightblue', foreground='black')
    #     # self.text_rus.tag_config('green',background='lightgreen',foreground='black')

class Fonts(object):
    def __init__(self,graph_):
        self.graph_=graph_
        self.conn=sqlite3.connect('V1.db')
        self.c=self.conn.cursor()
        self.selected=''



    def bold_(self):
        idx=self.graph_.word_.rnd_for_continue
        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag before: ',tag_)
        try:
            selected=self.graph_.text_rus.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='rus'
            self.graph_.text_rus.tag_add('bold',SEL_FIRST,SEL_LAST)
        except:
            pass
        try:
            selected=self.graph_.text_eng.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='eng'
            self.graph_.text_eng.tag_add('bold',SEL_FIRST,SEL_LAST)
        except:
            pass
        try:
            selected=self.graph_.text_esp.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='esp'
            self.graph_.text_esp.tag_add('bold',SEL_FIRST,SEL_LAST)
        except:
            pass

        if len(tag_)>0:
            new_tag_=tag_[0]+','+'%s:'%(text_w)+'bold:%s'%(selected)+','
        else:
            new_tag_=','+'%s:'%(text_w)+'bold:%s'%(selected)+','

        self.c.execute("UPDATE voc_table SET tag=? WHERE rowid=?",(new_tag_,idx))
        self.conn.commit()

        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()

        self.graph_.text_rus.tag_config('bold',font=('Arial',10,'bold'))
        self.graph_.text_eng.tag_config('bold',font=('Arial',10,'bold'))
        self.graph_.text_esp.tag_config('bold',font=('Arial',10,'bold'))

        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag after: ',tag_)




    def italic_(self):
        idx=self.graph_.word_.rnd_for_continue
        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag before: ',tag_)
        try:
            selected=self.graph_.text_rus.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='rus'
            self.graph_.text_rus.tag_add('italic',SEL_FIRST,SEL_LAST)
        except:
            pass
        try:
            selected=self.graph_.text_eng.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='eng'
            self.graph_.text_eng.tag_add('italic',SEL_FIRST,SEL_LAST)
        except:
            pass
        try:
            selected=self.graph_.text_esp.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='esp'
            self.graph_.text_esp.tag_add('italic',SEL_FIRST,SEL_LAST)
        except:
            pass

        if len(tag_)>0:
            new_tag_=tag_[0]+','+'%s:'%(text_w)+'italic:%s'%(selected)+','
        else:
            new_tag_=','+'%s:'%(text_w)+'italic:%s'%(selected)+','

        self.c.execute("UPDATE voc_table SET tag=? WHERE rowid=?",(new_tag_,idx))
        self.conn.commit()

        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag after: ',tag_)



        self.graph_.text_rus.tag_config('italic',font=('Arial',10,'italic','bold'))
        self.graph_.text_eng.tag_config('italic',font=('Arial',10,'italic','bold'))
        self.graph_.text_esp.tag_config('italic',font=('Arial',10,'italic','bold'))

    def blue_bg(self):
        idx=self.graph_.word_.rnd_for_continue
        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag before: ',tag_)
        try:
            selected=self.graph_.text_rus.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='rus'
            self.graph_.text_rus.tag_add('blue_bg',SEL_FIRST,SEL_LAST)
        except:
            pass
        try:
            selected=self.graph_.text_eng.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='eng'
            self.graph_.text_eng.tag_add('blue_bg',SEL_FIRST,SEL_LAST)
        except:
            pass
        try:
            selected=self.graph_.text_esp.get(SEL_FIRST,SEL_LAST)
            if len(selected)>1:
                text_w='esp'
            self.graph_.text_esp.tag_add('blue_bg',SEL_FIRST,SEL_LAST)
        except:
            pass
        if len(tag_)>0:
            new_tag_=tag_[0]+','+'%s:'%(text_w)+'blue_bg:%s'%(selected)+','
        else:
            new_tag_=','+'%s:'%(text_w)+'blue_bg:%s'%(selected)+','

        self.c.execute("UPDATE voc_table SET tag=? WHERE rowid=?",(new_tag_,idx))
        self.conn.commit()

        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()


        self.graph_.text_rus.tag_config('blue_bg',font='bold',background='#46a3ff')
        self.graph_.text_eng.tag_config('blue_bg',font='bold',background='#46a3ff')
        self.graph_.text_esp.tag_config('blue_bg',font='bold',background='#46a3ff')

        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag after: ',tag_)

    def clear_tag(self):
        try:
            selected=self.graph_.text_rus.get(SEL_FIRST,SEL_LAST)
            self.graph_.text_rus.tag_add('blank',SEL_FIRST,SEL_LAST)
            self.graph_.text_rus.tag_config('blank',background='old lace',font=())
        except:
            pass
        try:
            selected=self.graph_.text_eng.get(SEL_FIRST,SEL_LAST)
            self.graph_.text_eng.tag_add('blank',SEL_FIRST,SEL_LAST)
            self.graph_.text_eng.tag_config('blank',background='old lace',font=())
        except:
            pass
        try:
            selected=self.graph_.text_esp.get(SEL_FIRST,SEL_LAST)
            self.graph_.text_esp.tag_add('blank',SEL_FIRST,SEL_LAST)
            self.graph_.text_esp.tag_config('blank',background='old lace',font=())
        except:
            pass
        idx=self.graph_.word_.rnd_for_continue
        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag before: ', tag_)
        tag_=str(tag_)
        tag_=tag_.split(',')
        new_tag=''
        for w in tag_:
            if 'rus' in w or 'eng' in w or 'esp' in w:
                temp=(w.split(':'))[2]
                temp=temp.strip('\'')
                temp=temp.strip('\"')
                if selected not in w and temp not in selected:
                    print('selected: ', selected)
                    print('w2:  ', (w.split(':'))[2])
                    new_tag=new_tag+','+w


        self.c.execute("UPDATE voc_table SET tag=? WHERE rowid=?",(new_tag,idx))
        self.conn.commit()

        self.c.execute("SELECT tag FROM voc_table WHERE rowid=?",(idx,))
        tag_=self.c.fetchone()
        print('tag after: ', tag_)

class Delete_row(object):
    def __init__(self,graph_):
        self.graph_=graph_
        self.conn=sqlite3.connect('V1.db')
        self.c=self.conn.cursor()
        self.toDelete_rus=[]
        self.toDelete_eng=[]
        self.toDelete_esp=[]


    def toDelete(self):
        self.c.execute("SELECT * FROM voc_table")
        try:
            selected=self.graph_.text_rus.get(SEL_FIRST,SEL_LAST)
        except:
            pass

        try:
            selected=self.graph_.text_eng.get(SEL_FIRST,SEL_LAST)
        except:
            pass

        try:
            selected=self.graph_.text_esp.get(SEL_FIRST,SEL_LAST)
        except:
            pass

        for w in self.c.fetchall():
                if selected in w[1] or selected in w[2] or selected in w[3]:
                    print(w[0])
                    self.c.execute("DELETE FROM voc_table WHERE id_=?",(w[0],))
        self.conn.commit()



class Edit_(object):
    def __init__(self, graph_):
        self.graph_ = graph_

    def generic_func(self):
        self.edit_rus()
        self.edit_eng()
        self.edit_esp()

    def edit_rus(self):
        to_modif_rus = (self.graph_.text_rus.get('1.0', 'end'))
        if self.graph_.word_.rnd_for_continue == self.graph_.word_.total_count[self.graph_.word_.ind_backwards]:
            idx = (self.graph_.word_.rnd_for_continue)
            print('rus if: ', idx)
        else:
            idx = self.graph_.word_.total_count[self.graph_.word_.ind_backwards + 1]
            print('rus else: ', idx)
        conn = sqlite3.connect('V1.db')
        c = conn.cursor()
        c.execute("UPDATE voc_table SET rus_=? WHERE rowid=?", (to_modif_rus, idx))
        conn.commit()

    def edit_eng(self):
        to_modif_eng = (self.graph_.text_eng.get('1.0', 'end'))
        if self.graph_.word_.rnd_for_continue == self.graph_.word_.total_count[self.graph_.word_.ind_backwards]:
            idx = (self.graph_.word_.rnd_for_continue)
            print('eng if: ', idx)
        else:
            idx = self.graph_.word_.total_count[self.graph_.word_.ind_backwards + 1]
            print('eng else: ', idx)
        conn = sqlite3.connect('V1.db')
        c = conn.cursor()
        c.execute("UPDATE voc_table SET eng_=? WHERE rowid=?", (to_modif_eng, idx))
        conn.commit()

    def edit_esp(self):
        to_modif_esp = (self.graph_.text_esp.get('1.0', 'end'))
        if self.graph_.word_.rnd_for_continue == self.graph_.word_.total_count[self.graph_.word_.ind_backwards]:
            idx = (self.graph_.word_.rnd_for_continue)
            print('esp if: ', idx)
        else:
            idx = self.graph_.word_.total_count[self.graph_.word_.ind_backwards + 1]
            print('esp else: ', idx)

        conn = sqlite3.connect('V1.db')
        c = conn.cursor()
        c.execute("UPDATE voc_table SET esp_=? WHERE rowid=?", (to_modif_esp, idx))
        conn.commit()

    def edit_word_file(self):
        os.startfile('C:/C/Vocabulario_Completo_copy.docx')
        time.sleep(2)
        pg.moveTo(500, 400)
        pg.click(500, 400)
        pg.hotkey('ctrl', 'f')
        try:
            selected = self.graph_.text_eng.get(SEL_FIRST, SEL_LAST)
        except:
            None
        try:
            selected = self.graph_.text_esp.get(SEL_FIRST, SEL_LAST)
        except:
            None
        pg.typewrite(selected)
        pg.hotkey('enter')


class Nuevo(object):
    def __init__(self, graph_):
        self.graph_ = graph_
        self.conn=sqlite3.connect('V1.db')
        self.c=self.conn.cursor()


    def clear(self):
        self.graph_.text_rus.delete('1.0', 'end')
        self.graph_.text_eng.delete('1.0', 'end')
        self.graph_.text_esp.delete('1.0', 'end')

    def crear_(self):
        self.c.execute("SELECT COUNT(*) FROM voc_table")
        idx_raw=self.c.fetchall()
        idx=idx_raw[0][0]
        rus_esti_to_insert=self.graph_.text_rus.get('1.0','end')
        eng_to_insert=self.graph_.text_eng.get('1.0','end')
        esp_to_insert=self.graph_.text_esp.get('1.0','end')

        self.c.execute('INSERT INTO voc_table VALUES (?,?,?,?,?)',(idx+1,rus_esti_to_insert,eng_to_insert,esp_to_insert,''))
        self.conn.commit()



class Favorits(object):
    def __init__(self, graph_):
        start_fav=time.process_time()
        self.fav_list = []
        self.graph_ = graph_
        self.conn3 = sqlite3.connect('V2.db')
        #_____________________________________________
        self.c3 = self.conn3.cursor()
        self.c3.execute("SELECT * FROM fav")
        for w in self.c3.fetchall():
            self.fav_list.append(w)
        stop_fav=time.process_time()
        print('Fav time: %f'%(stop_fav-start_fav))

    def add_(self):

        to_add_rus = self.graph_.text_rus.get('1.0', 'end')
        to_add_eng = self.graph_.text_eng.get('1.0', 'end')
        to_add_esp = self.graph_.text_esp.get('1.0', 'end')

        self.fav_list.append((len(self.fav_list), to_add_rus, to_add_eng, to_add_eng))

    def rem_(self):
        idx = (self.graph_.word_.f_0_to_remove)
        for w in self.fav_list:
            if idx == w:
                self.fav_list.remove(w)

    def on_closing(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.c3.execute("DELETE FROM fav")
            for w in self.fav_list:
                self.c3.execute("INSERT INTO fav VALUES(?,?,?,?)", (w[0], w[1], w[2], w[3]))
            self.conn3.commit()
            self.graph_.master.destroy()


class Eesti(object):
    def __init__(self, graph_):
        self.est_count=1
        start_eesti=time.process_time()
        self.graph_ = graph_

        self.conn_eesti = sqlite3.connect('Eesti_base.db')
        self.c_eesti = self.conn_eesti.cursor()

        def to_update_estonian():
            self.c_eesti.execute("CREATE TABLE eesti_words(index_,est_,eng_,ex_)")
            wordDoc = Document('Vocabulario_Eesti.docx')
            idx = 0
            for table in wordDoc.tables:
                for row in table.rows:
                    list_ = []

                    for cell in row.cells:
                        temp_ = ''
                        for ch in cell.text:
                            if ord(ch) in range(65535):
                                temp_ = temp_ + ch
                        try:
                            list_.append(temp_)
                        except:
                            None
                    if len(list_) == 3:
                        self.c_eesti.execute("INSERT INTO eesti_words VALUES (?,?,?,?)", (idx, list_[0], list_[1], list_[2]))
                        idx += 1
                    elif len(list_) == 2:
                        self.c_eesti.execute("INSERT INTO eesti_words VALUES (?,?,?,?)", (idx, list_[0], list_[1], 'None'))
                        idx += 1
                    self.conn_eesti.commit()
        # to_update_estonian() #first delete Eesti_base.db

        stop_eesti=time.process_time()
        print('Eesti time: %f'%(stop_eesti- start_eesti))

    def getEestiWord(self):
        self.c_eesti.execute("SELECT Count(*) FROM eesti_words")
        eesti_lim = self.c_eesti.fetchall()
        print('eesti limit: ', eesti_lim)
        # rnd = random.randint(1, eesti_lim[0][0]) #Random estonian words
        rnd=self.est_count
        self.graph_.text_rus.delete('1.0', 'end')
        self.graph_.text_esp.delete('1.0', 'end')
        self.graph_.text_eng.delete('1.0', 'end')
        self.c_eesti.execute("SELECT * FROM eesti_words WHERE rowid=%d" % rnd)
        f = (self.c_eesti.fetchone())
        self.graph_.text_rus.insert('1.0', f[1])
        self.graph_.text_eng.insert('1.0', f[2])
        self.graph_.text_esp.insert('1.0', f[3])
        self.est_count+=1
        #self.word_.bg_colour()


class W0rd(object):
    def __init__(self, graph_):
        start_Word=time.process_time()
        self.rnd_for_continue = 0
        self.ind_backwards = -1
        self.rnd_fav = 0
        self.f_0_to_remove = []
        self.graph_ = graph_
        self.total_count = []
        #self.eesti_ = Eesti(self)
        self.conn = sqlite3.connect(r'V1.db')
        self.conn_frases = sqlite3.connect(':memory:')

        self.conn_favorits = sqlite3.connect('V2.db')
        #_____________________________________________________
        self.c = self.conn.cursor()
        self.c_frases = self.conn_frases.cursor()

        self.c_frases.execute("CREATE TABLE frases_words(id_,rus_f,eng_f)")

        self.c_favorits = self.conn_favorits.cursor()
        self.search_list_rus = []
        self.search_list_eng = []
        self.search_list_esp = []
        self.c.execute("SELECT *FROM voc_table")
        for w in self.c.fetchall():
            if 'Frases' in w[1]:
                self.c_frases.execute("INSERT INTO frases_words VALUES(?,?,?)", (w[0], w[1], w[2]))
                self.conn_frases.commit()
            if len(w) == 5:
                self.search_list_rus.append(w[1])
                self.search_list_eng.append(w[2])
                self.search_list_esp.append(w[3])
            elif len(w) == 4:
                self.search_list_rus.append(w[1])
                self.search_list_eng.append(w[2])
                self.search_list_esp.append('None')
        stop_Word=time.process_time()
        print('Word time: %f'%(stop_Word- start_Word))


    # def bg_colour(self):
    #     if self.graph_.text_rus.search('ʻ', '1.0', 'end'):
    #         open_ = self.graph_.text_rus.search('ʻ', '1.0', 'end')
    #         close_ = self.graph_.text_rus.search('ʼ', '1.0', 'end')
    #         if open_ and close_:

    #             self.graph_.text_rus.tag_add('blue', open_, close_)
    #             self.graph_.text_rus.tag_config('blue', background='lightblue')

    #     elif self.graph_.text_eng.search('ʻ', '1.0', 'end'):
    #         open_ = self.graph_.text_eng.search('ʻ', '1.0', 'end')
    #         close_ = self.graph_.text_eng.search('ʼ', '1.0', 'end')
    #         if open_ and close_:

    #             self.graph_.text_eng.tag_add('blue', open_, close_)
    #             self.graph_.text_eng.tag_config('blue', background='lightblue')

    #     elif self.graph_.text_esp.search('ʻ', '1.0', 'end'):
    #         open_ = self.graph_.text_esp.search('ʻ', '1.0', 'end')
    #         close_ = self.graph_.text_esp.search('ʼ', '1.0', 'end')
    #         if open_ and close_:

    #             self.graph_.text_esp.tag_add('blue', open_, close_)
    #             self.graph_.text_esp.tag_config('blue', background='lightblue')

    def getWord(self):

        if (self.graph_.sel.get()) == 'True':

            self.c_frases.execute("SELECT Count(*) FROM frases_words")
            frase_lim = self.c_frases.fetchall()
            rnd = random.randint(1, frase_lim[0][0])
            self.graph_.text_rus.delete('1.0', 'end')
            self.graph_.text_esp.delete('1.0', 'end')
            self.graph_.text_eng.delete('1.0', 'end')
            self.c_frases.execute("SELECT * FROM frases_words WHERE rowid=%d" % rnd)
            f = (self.c_frases.fetchone())
            self.graph_.text_rus.insert('1.0', f[1])

            #self.bg_colour()

        elif (self.graph_. eesti_sel .get()) == 'True':
            self.graph_.eesti_.getEestiWord()

        elif (self.graph_.fav_sel.get()) == 'True':

            rnd = random.randint(0, len(self.graph_.f_word.fav_list) - 1)
            self.f_0_to_remove = self.graph_.f_word.fav_list[rnd]

            self.graph_.text_rus.delete('1.0', 'end')
            self.graph_.text_esp.delete('1.0', 'end')
            self.graph_.text_eng.delete('1.0', 'end')
            if len(self.graph_.f_word.fav_list[rnd]) == 4:
                self.graph_.text_rus.insert('1.0', self.graph_.f_word.fav_list[rnd][1])
                self.graph_.text_eng.insert('1.0', self.graph_.f_word.fav_list[rnd][2])
                self.graph_.text_esp.insert('1.0', self.graph_.f_word.fav_list[rnd][3])
            elif len(self.graph_.f_word.fav_list[rnd]) == 3:
                self.graph_.text_rus.insert('1.0', self.graph_.f_word.fav_list[rnd][1])
                self.graph_.text_eng.insert('1.0', self.graph_.f_word.fav_list[rnd][2])
                self.graph_.text_esp.insert('1.0', 'None')
            #self.bg_colour()

        elif (self.graph_.goBack.get()) == 'True':
            self.ind_backwards = self.ind_backwards - 1
            self.c.execute("select * from voc_table where rowid=%d" % self.total_count[self.ind_backwards])

            print('back: ', self.total_count[self.ind_backwards])
            temp_list = self.c.fetchone()
            try:
                self.graph_.text_rus.delete('1.0', 'end')
                self.graph_.text_rus.insert('1.0', temp_list[1])

                self.graph_.text_eng.delete('1.0', 'end')
                self.graph_.text_eng.insert('1.0', temp_list[2])

                self.graph_.text_esp.delete('1.0', 'end')
                self.graph_.text_esp.insert('1.0', temp_list[3])
            except:
                None
            #self.bg_colour()

        else:
            self.ind_backwards = -1
            self.c.execute("SELECT Count(*) FROM voc_table")
            lim_ = self.c.fetchall()
            self.rnd_for_continue = random.randint(1, lim_[0][0])
            print(lim_[0][0], ' - ', self.rnd_for_continue)
            self.graph_.label_info2.config(text='  %d - %d' % (lim_[0][0], self.rnd_for_continue))
            self.total_count.append(self.rnd_for_continue)
            print('straight: ', self.total_count)
            self.c.execute("SELECT * FROM voc_table WHERE rowid=%d" % self.rnd_for_continue)
            temp_list = (self.c.fetchone())

            getTag=temp_list[-1]
            self.graph_.text_rus.delete('1.0', 'end')
            self.graph_.text_rus.insert('1.0', temp_list[1])
            self.graph_.text_eng.delete('1.0', 'end')
            self.graph_.text_eng.insert('1.0', temp_list[2])
            self.graph_.text_esp.delete('1.0', 'end')
            self.graph_.text_esp.insert('1.0', temp_list[3])

            self.print_fonts(getTag)
            #self.bg_colour()

    def print_fonts(self, getTag):
        countVar=StringVar()
        if getTag:
            getTag=getTag.split(',')
            for w in getTag:
                if 'bold' in w:
                    w=w.split(':')
                    #print('w: ', w[0],w[1],w[2])
                    if w[0]=='rus':
                        start_pos=self.graph_.text_rus.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_rus.tag_add('bold',start_pos,end_pos)
                    elif w[0]=='eng':
                        start_pos=self.graph_.text_eng.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_eng.tag_add('bold',start_pos,end_pos)
                    elif w[0]=='esp':
                        start_pos=self.graph_.text_esp.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_esp.tag_add('bold',start_pos,end_pos)


                elif 'italic' in w:
                    w=w.split(':')
                    print('w: ', w[0],w[1],w[2])
                    if w[0]=='rus':
                        start_pos=self.graph_.text_rus.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_rus.tag_add('italic',start_pos,end_pos)
                    elif w[0]=='eng':
                        start_pos=self.graph_.text_eng.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_eng.tag_add('italic',start_pos,end_pos)
                    elif w[0]=='esp':
                        start_pos=self.graph_.text_esp.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_esp.tag_add('italic',start_pos,end_pos)


                elif 'blue_bg' in w:
                    w=w.split(':')
                    print('w: ', w[0],w[1],w[2])
                    if w[0]=='rus':
                        start_pos=self.graph_.text_rus.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_rus.tag_add('blue_bg',start_pos,end_pos)
                    elif w[0]=='eng':
                        start_pos=self.graph_.text_eng.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_eng.tag_add('blue_bg',start_pos,end_pos)
                    elif w[0]=='esp':
                        start_pos=self.graph_.text_esp.search(w[2],'1.0',stopindex=END,count=countVar)
                        end_pos='%s+%sc'%(start_pos,countVar.get())
                        self.graph_.text_esp.tag_add('blue_bg',start_pos,end_pos)


        self.graph_.text_rus.tag_config('bold',font=('Arial',10,'bold'))
        self.graph_.text_eng.tag_config('bold',font=('Arial',10,'bold'))
        self.graph_.text_esp.tag_config('bold',font=('Arial',10,'bold'))

        self.graph_.text_rus.tag_config('italic',font=('Arial',10,'italic','bold'))
        self.graph_.text_eng.tag_config('italic',font=('Arial',10,'italic','bold'))
        self.graph_.text_esp.tag_config('italic',font=('Arial',10,'italic','bold'))


        self.graph_.text_rus.tag_config('blue_bg',font='bold',background='#46a3ff')
        self.graph_.text_eng.tag_config('blue_bg',font='bold',background='#46a3ff')
        self.graph_.text_esp.tag_config('blue_bg',font='bold',background='#46a3ff')



def main():
    location_V1 = 'V1.db'
    location_V3 = 'V3.db'
    location_Voc_Compl = 'Vocabulario_Completo_Ruso_Eng_Esp_Eesti.docx'
    create_V3_combine_V3_with_V1 = Create_check_db()

    root = Tk()

    root.geometry('700x450+200+100')
    root.resizable(False, False)
    new_play = Graph(root)

    root.mainloop()


main()
